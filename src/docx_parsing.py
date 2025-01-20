import base64
from io import BytesIO

from docx import Document

from src.models import FileImage, FileText


def docx_extract_texts_and_images(file_content: bytes):
    """
    Parse a .docx file to extract texts, images, Markdown tables, and drawings,
    preserving the sequence of text and tables.

    Args:
        file_content (bytes): Bytes of the file.

    Returns:
        dict: A dictionary with extracted texts, images, Markdown tables, and drawings.
    """
    doc = Document(BytesIO(file_content))
    markdown_content = ""
    images = []
    markdown_tables = []
    drawings = []

    # Keep track of table index
    table_index = 0

    # Iterate through all block items in the document
    for block in doc.element.body.iter():
        if block.tag.endswith("p"):  # Paragraph
            # Find the corresponding paragraph object
            for paragraph in doc.paragraphs:
                if paragraph._element == block and paragraph.text.strip():
                    markdown_content += f"\n{paragraph.text.strip()}"
                    break

        elif block.tag.endswith("tbl"):  # Table
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                table_data = []
                max_cols = max(len(row.cells) for row in table.rows)

                # Extract table rows
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.replace("\n", "<br>").strip())
                    table_data.append(row_data)

                # Convert table to Markdown
                if table_data:
                    table_markdown = []
                    header = table_data[0]
                    table_markdown.append("| " + " | ".join(header) + " |")
                    table_markdown.append("| " + " | ".join(["---"] * max_cols) + " |")
                    for row in table_data[1:]:
                        table_markdown.append("| " + " | ".join(row) + " |")

                    # Add to both content and separate tables list
                    table_str = "\n".join(table_markdown)
                    markdown_content += f"\n{table_str}"
                    markdown_tables.append(table_str)

                table_index += 1

    # Extract images and drawings
    image_counter = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_base64 = base64.b64encode(image_data).decode("utf-8")
            images.append(
                FileImage(
                    page_no=0,  # Assuming rendering is dynamic
                    image_no=image_counter,
                    image_base64=image_base64,
                )
            )
            image_counter += 1
        elif "drawing" in rel.target_ref:
            drawings.append(rel.target_ref)

    return {
        "texts": [
            FileText(text=markdown_content, page_no=0)
        ],  # Sequential text and tables
        "images": images,
        # "tables": markdown_tables,  # Separate list of tables
        # "drawings": drawings,
    }
